import os
import re
from datetime import datetime

from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.hashers import make_password
from django.contrib.auth.mixins import LoginRequiredMixin, UserPassesTestMixin
from django.http import FileResponse
from django.conf import settings
from django.db.models import Q, Count
from django.views import View
from django.views.generic import TemplateView, ListView, FormView, DetailView, UpdateView
from django.urls import reverse_lazy


from docx import Document
from docx.shared import RGBColor
from openpyxl import Workbook

from .models import (
    User, Role, Graduate, Employer, Employment, EmploymentStatus,
    Feedback, Report, RegistrationRequest
)


# ========================
# МИКСИНЫ
# ========================

class RoleRequiredMixin(UserPassesTestMixin):
    """Миксин для проверки роли пользователя"""
    allowed_roles = []

    def test_func(self):
        if not self.request.user.is_authenticated:
            return False
        if hasattr(self.request.user, 'role') and self.request.user.role:
            return self.request.user.role.name in self.allowed_roles
        return False

    def handle_no_permission(self):
        messages.error(self.request, "У вас недостаточно прав для доступа к этой странице.")
        return redirect('muiv_graduation_system:index')


# ========================
# ОСНОВНЫЕ СТРАНИЦЫ
# ========================

class IndexView(TemplateView):
    """Главная страница"""
    template_name = 'index.html'


class AboutView(TemplateView):
    """О проекте"""
    template_name = 'about.html'


class ContactsView(TemplateView):
    """Контакты"""
    template_name = 'contacts.html'



class FeedbackView(View):
    """Обратная связь"""
    template_name = 'feedback.html'

    def get(self, request):
        return render(request, self.template_name)

    def post(self, request):
        subject = request.POST.get('subject', 'Без темы')
        message = request.POST.get('message', '').strip()

        if not message:
            messages.error(request, "Сообщение не может быть пустым")
            return render(request, self.template_name)

        # Если пользователь авторизован — используем его
        if request.user.is_authenticated:
            user = request.user
            Feedback.objects.create(user=user, subject=subject, message=message)
            messages.success(request, "Ваше сообщение отправлено! Спасибо за обратную связь.")
            return redirect('muiv_graduation_system:index')

        # Пользователь НЕ авторизован — обрабатываем как гость
        email = request.POST.get('email', '').strip()

        # Валидация email
        if not email:
            messages.error(request, "Пожалуйста, укажите email.")
            return render(request, self.template_name)

        if not re.match(r"[^@]+@[^@]+\.[^@]+", email):
            messages.error(request, "Некорректный адрес email.")
            return render(request, self.template_name)

        # 🔍 Проверяем, существует ли уже пользователь с таким email
        if User.objects.filter(email=email).exists():
            messages.error(
                request,
                "На этот email уже зарегистрирован аккаунт. Вам необходимо войти в систему."
            )
            return render(request, self.template_name)

        # Если email свободен — создаём временного пользователя
        try:
            role, _ = Role.objects.get_or_create(name='graduate')
            username = f"guest_{datetime.now().strftime('%Y%m%d%H%M%S%f')}"
            user = User.objects.create(
                username=username,
                email=email,
                password=make_password('guest123'),
                role=role
            )
        except Exception as e:
            # На случай, если что-то пошло не так при создании (например, дубль username)
            messages.error(request, "Не удалось создать временного пользователя. Попробуйте позже.")
            return render(request, self.template_name)

        # Сохраняем обратную связь
        Feedback.objects.create(user=user, subject=subject, message=message)
        messages.success(request, "Ваше сообщение отправлено! Спасибо за обратную связь.")
        return redirect('muiv_graduation_system:index')


# ========================
# АВТОРИЗАЦИЯ
# ========================

class LoginView(View):
    """Вход в систему"""
    template_name = 'login.html'

    def get(self, request):
        if request.user.is_authenticated:
            return redirect('muiv_graduation_system:profile')
        return render(request, self.template_name)

    def post(self, request):
        username = request.POST.get('username', '')
        password = request.POST.get('password', '')

        user = authenticate(request, username=username, password=password)

        if user is not None:
            login(request, user)
            display_name = user.username

            # Получаем полное имя для выпускника
            if user.role and user.role.name == 'graduate':
                try:
                    grad = Graduate.objects.get(user=user)
                    display_name = grad.full_name
                except Graduate.DoesNotExist:
                    pass

            messages.success(request, f"Добро пожаловать, {display_name}!")
            return redirect('muiv_graduation_system:profile')
        else:
            messages.error(request, "Неверный логин или пароль")
            return render(request, self.template_name)


class LogoutView(View):
    """Выход из системы"""

    def get(self, request):
        logout(request)
        messages.info(request, "Вы вышли из системы")
        return redirect('muiv_graduation_system:index')


class RegisterView(TemplateView):
    """Выбор роли при регистрации"""
    template_name = 'register.html'


class RegisterAsView(View):
    """Регистрация с выбранной ролью"""

    def get(self, request, role):
        if role == 'graduate':
            return render(request, 'register_graduate.html')
        elif role == 'manager':
            return render(request, 'register_manager.html')
        else:
            messages.error(request, "Недопустимая роль")
            return redirect('muiv_graduation_system:register')

    def post(self, request, role):
        if role == 'graduate':
            return self._register_graduate(request)
        elif role == 'manager':
            return self._register_manager(request)
        else:
            messages.error(request, "Недопустимая роль")
            return redirect('muiv_graduation_system:register')

    def _register_graduate(self, request):
        """Регистрация выпускника"""
        last_name = request.POST.get('last_name', '').strip()
        first_name = request.POST.get('first_name', '').strip()
        middle_name = request.POST.get('middle_name', '').strip()
        grad_year = request.POST.get('graduation_year', '')
        email = request.POST.get('email', '').strip()
        username = request.POST.get('username', '').strip()
        password = request.POST.get('password', '')

        # Валидация
        if not last_name or not first_name:
            messages.error(request, "Фамилия и имя обязательны")
            return render(request, 'register_graduate.html')

        if not re.match(r"[^@]+@[^@]+\.[^@]+", email):
            messages.error(request, "Некорректный email")
            return render(request, 'register_graduate.html')

        if User.objects.filter(username=username).exists():
            messages.error(request, "Логин занят")
            return render(request, 'register_graduate.html')

        if User.objects.filter(email=email).exists():
            messages.error(request, "Email уже используется")
            return render(request, 'register_graduate.html')

        # Создание пользователя
        full_name = f"{last_name} {first_name}" + (f" {middle_name}" if middle_name else "")
        role_obj, _ = Role.objects.get_or_create(name='graduate')

        user = User.objects.create(
            username=username,
            email=email,
            password=make_password(password),
            role=role_obj
        )

        Graduate.objects.create(
            user=user,
            full_name=full_name,
            graduation_year=int(grad_year),
            email=email
        )

        messages.success(request, "Регистрация завершена! Войдите в систему.")
        return redirect('muiv_graduation_system:login')

    def _register_manager(self, request):
        """Регистрация менеджера (заявка)"""
        username = request.POST.get('username', '').strip()
        email = request.POST.get('email', '').strip()
        password = request.POST.get('password', '')
        full_name = request.POST.get('full_name', '').strip()

        # Проверка уникальности
        if (User.objects.filter(Q(username=username) | Q(email=email)).exists() or
                RegistrationRequest.objects.filter(Q(username=username) | Q(email=email)).exists()):
            messages.error(request, "Логин или email уже используется")
            return render(request, 'register_manager.html')

        # Создание заявки
        RegistrationRequest.objects.create(
            username=username,
            email=email,
            password_hash=make_password(password),
            full_name=full_name
        )

        messages.info(request, "Заявка отправлена. Ожидайте одобрения администратором.")
        return redirect('muiv_graduation_system:index')


# ========================
# ПРОФИЛЬ
# ========================

class ProfileView(LoginRequiredMixin, View):
    """Профиль пользователя (роутер по ролям)"""
    login_url = 'muiv_graduation_system:login'

    def get(self, request):
        user = request.user

        if not hasattr(user, 'role') or not user.role:
            messages.error(request, "У пользователя не назначена роль")
            return redirect('muiv_graduation_system:index')

        role_name = user.role.name

        if role_name == 'graduate':
            return self._graduate_profile(request)
        elif role_name == 'manager':
            return self._manager_profile(request)
        elif role_name == 'admin':
            return self._admin_profile(request)

        return redirect('muiv_graduation_system:index')

    def _graduate_profile(self, request):
        """Профиль выпускника"""
        try:
            graduate = Graduate.objects.select_related('employment__employer', 'employment__status').get(
                user=request.user)
        except Graduate.DoesNotExist:
            messages.error(request, "Профиль выпускника не найден")
            return redirect('muiv_graduation_system:index')

        return render(request, 'profile/graduate.html', {'graduate': graduate})

    def _manager_profile(self, request):
        """Профиль менеджера"""
        return render(request, 'profile/manager.html', {'user': request.user})

    def _admin_profile(self, request):
        """Профиль администратора"""
        stats = {
            'total_users': User.objects.count(),
            'total_graduates': Graduate.objects.count(),
            'employed_graduates': Employment.objects.filter(status__name='трудоустроен').count()
        }
        return render(request, 'profile/admin.html', {'user': request.user, 'stats': stats})


# ========================
# ВЫПУСКНИК: РЕДАКТИРОВАНИЕ И ЭКСПОРТ
# ========================

class EditGraduateView(RoleRequiredMixin, View):
    """Редактирование профиля выпускника (сам)"""
    allowed_roles = ['graduate']
    template_name = 'graduate/edit.html'

    def get(self, request):
        graduate = get_object_or_404(Graduate, user=request.user)
        employment = getattr(graduate, 'employment', None)
        statuses = EmploymentStatus.objects.all()
        employer_name = employment.employer.name if employment and employment.employer else ''

        return render(request, self.template_name, {
            'graduate': graduate,
            'employment': employment,
            'statuses': statuses,
            'employer_name': employer_name,
            'errors': {},
            'form_data': {},
        })

    def post(self, request):
        graduate = get_object_or_404(Graduate, user=request.user)
        employment = getattr(graduate, 'employment', None)

        # === Сбор данных из формы ===
        full_name = request.POST.get('full_name', '').strip()
        graduation_year_str = request.POST.get('graduation_year', '').strip()
        email = request.POST.get('email', '').strip()
        faculty = request.POST.get('faculty', '').strip()
        specialization = request.POST.get('specialization', '').strip()
        phone = request.POST.get('phone', '').strip()

        errors = {}
        missing_fields = []

        # === Валидация обязательных полей ===
        if not full_name:
            errors['full_name'] = True
            missing_fields.append("ФИО")
        if not graduation_year_str:
            errors['graduation_year'] = True
            missing_fields.append("Год выпуска")
        elif not graduation_year_str.isdigit():
            errors['graduation_year'] = "Год выпуска должен быть числом."
        else:
            grad_year = int(graduation_year_str)
            if grad_year < 2000 or grad_year > 2030:
                errors['graduation_year'] = "Год должен быть между 2000 и 2030."

        if not email:
            errors['email'] = True
            missing_fields.append("Email")
        elif not re.match(r"[^@]+@[^@]+\.[^@]+", email):
            errors['email'] = "Некорректный формат email."
        else:
            if User.objects.filter(email=email).exclude(id=request.user.id).exists():
                errors['email'] = "Email уже используется другим пользователем."

        # Обязательные поля при редактировании
        if not faculty:
            errors['faculty'] = True
            missing_fields.append("Факультет")
        if not specialization:
            errors['specialization'] = True
            missing_fields.append("Специальность")
        if not phone:
            errors['phone'] = True
            missing_fields.append("Телефон")

        # === Обработка статуса трудоустройства (через ID!) ===
        status_id = request.POST.get('employment_status', '').strip()

        # Управление записью Employment
        if status_id:
            try:
                status = EmploymentStatus.objects.get(id=status_id)
                if not employment:
                    employment = Employment.objects.create(graduate=graduate, status=status)
                else:
                    employment.status = status
                    employment.save()
            except EmploymentStatus.DoesNotExist:
                messages.warning(request, "Выбранный статус трудоустройства не найден.")
                status = None
        else:
            # Статус снят
            if employment:
                employment.status = None
                employment.save()

        # === Показываем общую ошибку, если есть незаполненные обязательные поля ===
        if missing_fields:
            field_list = ", ".join(missing_fields)
            messages.error(request, f"❌ Заполните поля: {field_list}.")

        # === Если есть ошибки — не сохраняем, возвращаем форму ===
        if errors:
            statuses = EmploymentStatus.objects.all()
            employer_name = employment.employer.name if employment and employment.employer else ''
            return render(request, self.template_name, {
                'graduate': graduate,
                'employment': employment,
                'statuses': statuses,
                'employer_name': employer_name,
                'errors': errors,
                'form_data': request.POST,
            })

        # === Сохранение личных данных ===
        graduate.full_name = full_name
        graduate.graduation_year = int(graduation_year_str)
        graduate.faculty = faculty
        graduate.specialization = specialization
        graduate.email = email
        graduate.phone = phone
        graduate.save()

        # === Сохранение остальных данных трудоустройства (если Employment существует) ===
        if employment:
            employment.job_title = request.POST.get('job_title', '').strip() or None
            salary_str = request.POST.get('salary', '').strip()
            employment.salary = int(salary_str) if salary_str.isdigit() else None

            start_date_str = request.POST.get('start_date', '').strip()
            if start_date_str:
                try:
                    employment.start_date = datetime.strptime(start_date_str, '%Y-%m-%d').date()
                except ValueError:
                    employment.start_date = None
            else:
                employment.start_date = None

            employer_name = request.POST.get('employer_name', '').strip()
            if employer_name:
                employer, _ = Employer.objects.get_or_create(name=employer_name)
                employment.employer = employer
            else:
                employment.employer = None

            employment.save()

        messages.success(request, "Данные успешно обновлены")
        return redirect('muiv_graduation_system:profile')


class ExportMyDataView(RoleRequiredMixin, View):
    """Экспорт личных данных выпускника"""
    allowed_roles = ['graduate']

    def get(self, request):
        graduate = get_object_or_404(Graduate, user=request.user)
        employment = getattr(graduate, 'employment', None)

        doc = Document()

        # Заголовок документа
        title = doc.add_heading('Персональный отчёт о трудоустройстве', 0)
        title.alignment = 1  # Центрирование

        # Информация о выпускнике
        doc.add_paragraph(f'Выпускник: {graduate.full_name}', style='Heading 2')
        doc.add_paragraph(f'Дата формирования: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
        doc.add_paragraph('')  # Пустая строка

        # ========================
        # ЛИЧНЫЕ ДАННЫЕ
        # ========================
        doc.add_heading('📋 Личные данные', level=1)

        # Таблица с личными данными
        personal_table = doc.add_table(rows=6, cols=2)
        personal_table.style = 'Light Grid Accent 1'

        personal_data = [
            ('ФИО', graduate.full_name),
            ('Год выпуска', str(graduate.graduation_year)),
            ('Факультет', graduate.faculty or 'Не указан'),
            ('Специальность', graduate.specialization or 'Не указана'),
            ('Email', graduate.email),
            ('Телефон', graduate.phone or 'Не указан'),
        ]

        for i, (label, value) in enumerate(personal_data):
            row = personal_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            # Жирный шрифт для лейблов
            row.cells[0].paragraphs[0].runs[0].bold = True

        doc.add_paragraph('')  # Пустая строка

        # ========================
        # ТРУДОУСТРОЙСТВО
        # ========================
        doc.add_heading('💼 Информация о трудоустройстве', level=1)

        if employment:
            # Статус трудоустройства
            status_para = doc.add_paragraph()
            status_para.add_run('Статус: ').bold = True

            if employment.status:
                status_run = status_para.add_run(employment.status.name)
                status_run.font.size = 12
                # Цвет в зависимости от статуса
                if employment.status.name.lower() == 'трудоустроен':
                    status_run.font.color.rgb = RGBColor(34, 139, 34)  # Зелёный
                elif employment.status.name.lower() == 'в поиске':
                    status_run.font.color.rgb = RGBColor(255, 165, 0)  # Оранжевый
                else:
                    status_run.font.color.rgb = RGBColor(220, 20, 60)  # Красный
            else:
                status_para.add_run('Не указан')

            doc.add_paragraph('')

            # Таблица с данными о работе
            employment_table = doc.add_table(rows=5, cols=2)
            employment_table.style = 'Light Grid Accent 1'

            employment_data = [
                ('Работодатель', employment.employer.name if employment.employer else 'Не указан'),
                ('Должность', employment.job_title or 'Не указана'),
                ('Зарплата', f'{employment.salary:,} ₽/мес.' if employment.salary else 'Не указана'),
                ('Дата начала работы',
                 employment.start_date.strftime('%d.%m.%Y') if employment.start_date else 'Не указана'),
                ('Стаж работы',
                 self._calculate_work_experience(employment.start_date) if employment.start_date else '—'),
            ]

            for i, (label, value) in enumerate(employment_data):
                row = employment_table.rows[i]
                row.cells[0].text = label
                row.cells[1].text = str(value)
                row.cells[0].paragraphs[0].runs[0].bold = True
        else:
            # Если трудоустройство не указано
            no_employment = doc.add_paragraph()
            no_employment.add_run('ℹ️ Информация о трудоустройстве отсутствует').italic = True
            no_employment.alignment = 1  # Центрирование

            doc.add_paragraph('')
            doc.add_paragraph('Рекомендации:')
            recommendations = [
                'Обновите информацию о трудоустройстве в личном кабинете',
                'Обратитесь в отдел по трудоустройству для консультации',
                'Воспользуйтесь карьерными услугами университета'
            ]
            for rec in recommendations:
                doc.add_paragraph(f'  • {rec}', style='List Bullet')

        # ========================
        # ФУТЕР
        # ========================
        doc.add_paragraph('')
        doc.add_paragraph('_' * 60)
        footer = doc.add_paragraph(
            f'Документ сформирован автоматически системой учёта трудоустройства выпускников МУ им. С.Ю. Витте'
        )
        footer.alignment = 1
        footer.runs[0].font.size = 8
        footer.runs[0].italic = True

        # Сохранение файла
        filename = f"personal_report_{graduate.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(settings.BASE_DIR, 'static', 'reports', filename)
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        doc.save(filepath)

        return FileResponse(open(filepath, 'rb'), as_attachment=True, filename=filename)

    def _calculate_work_experience(self, start_date):
        """Вычисление стажа работы"""
        if not start_date:
            return '—'

        from datetime import date
        today = date.today()
        delta = today - start_date

        years = delta.days // 365
        months = (delta.days % 365) // 30
        days = (delta.days % 365) % 30

        parts = []
        if years > 0:
            parts.append(f'{years} {self._pluralize(years, "год", "года", "лет")}')
        if months > 0:
            parts.append(f'{months} {self._pluralize(months, "месяц", "месяца", "месяцев")}')
        if not parts and days > 0:
            parts.append(f'{days} {self._pluralize(days, "день", "дня", "дней")}')

        return ' '.join(parts) if parts else 'Менее месяца'

    def _pluralize(self, n, form1, form2, form5):
        """Склонение слов по числам"""
        n = abs(n) % 100
        if n >= 5 and n <= 20:
            return form5
        n = n % 10
        if n == 1:
            return form1
        if n >= 2 and n <= 4:
            return form2
        return form5


# ========================
# МЕНЕДЖЕР: УПРАВЛЕНИЕ ВЫПУСКНИКАМИ
# ========================

class ManagerGraduatesView(RoleRequiredMixin, ListView):
    """Список выпускников для менеджера"""
    allowed_roles = ['manager', 'admin']
    template_name = 'manager/graduates.html'
    context_object_name = 'graduates'
    paginate_by = 25

    def get_queryset(self):
        return Graduate.objects.select_related(
            'employment__employer',
            'employment__status'
        ).all()


class EditGraduateByManagerView(RoleRequiredMixin, View):
    """Редактирование выпускника менеджером"""
    allowed_roles = ['manager', 'admin']
    template_name = 'manager/edit_graduate.html'

    def get(self, request, grad_id):
        graduate = get_object_or_404(Graduate, id=grad_id)
        employment = getattr(graduate, 'employment', None)
        statuses = EmploymentStatus.objects.all()
        employer_name = employment.employer.name if employment and employment.employer else ''

        return render(request, self.template_name, {
            'graduate': graduate,
            'employment': employment,
            'statuses': statuses,
            'employer_name': employer_name
        })

    def post(self, request, grad_id):
        graduate = get_object_or_404(Graduate, id=grad_id)
        employment = getattr(graduate, 'employment', None)

        # Обновление данных (аналогично EditGraduateView)
        graduate.full_name = request.POST.get('full_name', '')
        graduate.graduation_year = int(request.POST.get('graduation_year', 0))
        graduate.faculty = request.POST.get('faculty', '') or None
        graduate.specialization = request.POST.get('specialization', '') or None
        graduate.phone = request.POST.get('phone', '') or None
        graduate.email = request.POST.get('email', '')
        graduate.save()

        status_name = request.POST.get('employment_status', '')
        if status_name:
            status, _ = EmploymentStatus.objects.get_or_create(name=status_name)
            if not employment:
                employment = Employment.objects.create(graduate=graduate, status=status)
            else:
                employment.status = status

        if employment:
            employment.job_title = request.POST.get('job_title', '') or None
            salary_str = request.POST.get('salary', '')
            employment.salary = int(salary_str) if salary_str else None

            start_date_str = request.POST.get('start_date', '')
            if start_date_str:
                try:
                    employment.start_date = datetime.strptime(start_date_str, '%Y-%m-%d').date()
                except ValueError:
                    employment.start_date = None

            employer_name = request.POST.get('employer_name', '')
            if employer_name:
                employer, _ = Employer.objects.get_or_create(name=employer_name)
                employment.employer = employer
            else:
                employment.employer = None

            employment.save()

        messages.success(request, "Данные обновлены")
        return redirect('muiv_graduation_system:manager_graduates')


# ========================
# ПОИСК И ЭКСПОРТ
# ========================

class SearchGraduatesView(RoleRequiredMixin, ListView):
    """Поиск выпускников"""
    allowed_roles = ['manager', 'admin']
    template_name = 'manager/search.html'
    context_object_name = 'graduates'
    paginate_by = 25

    def get_queryset(self):
        query = self.request.GET.get('query', '').strip()
        graduates = Graduate.objects.select_related(
            'employment__employer',
            'employment__status'
        )

        if query:
            graduates = graduates.filter(
                Q(full_name__icontains=query) |
                Q(specialization__icontains=query) |
                Q(faculty__icontains=query) |
                Q(employment__employer__name__icontains=query) |
                Q(employment__job_title__icontains=query)
            )

        return graduates.distinct()

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['query'] = self.request.GET.get('query', '')
        return context


class ExportSearchResultsView(RoleRequiredMixin, View):
    """Экспорт результатов поиска"""
    allowed_roles = ['manager', 'admin']

    def get(self, request, format):
        query = request.GET.get('query', '').strip()
        graduates = Graduate.objects.select_related('employment__employer', 'employment__status')

        if query:
            graduates = graduates.filter(
                Q(full_name__icontains=query) |
                Q(specialization__icontains=query) |
                Q(faculty__icontains=query) |
                Q(employment__employer__name__icontains=query) |
                Q(employment__job_title__icontains=query)
            ).distinct()
        else:
            graduates = graduates.all()

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"graduates_report_{timestamp}.{format}"
        filepath = os.path.join(settings.BASE_DIR, 'static', 'reports', filename)
        os.makedirs(os.path.dirname(filepath), exist_ok=True)

        if format == 'docx':
            self._export_docx(graduates, filepath, query)
        elif format == 'xlsx':
            self._export_xlsx(graduates, filepath, query)
        else:
            messages.error(request, "Неподдерживаемый формат отчёта")
            return redirect('muiv_graduation_system:search_graduates')

        # Сохранение в базу
        Report.objects.create(
            title=f"Отчёт выпускников от {datetime.now().strftime('%d.%m.%Y')}",
            generated_by=request.user,
            format=format,
            filepath=filepath
        )

        return FileResponse(open(filepath, 'rb'), as_attachment=True, filename=filename)

    def _export_docx(self, graduates, filepath, query):
        """Экспорт в DOCX с красивым форматированием"""
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Заголовок
        title = doc.add_heading('Отчёт по выпускникам', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Информация о поиске
        if query:
            search_info = doc.add_paragraph()
            search_info.add_run('🔍 Поисковый запрос: ').bold = True
            search_info.add_run(f'"{query}"')
            search_info.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Метаданные
        meta = doc.add_paragraph()
        meta.add_run(f'📅 Дата формирования: {datetime.now().strftime("%d.%m.%Y %H:%M")}\n')
        meta.add_run(f'👥 Найдено выпускников: {graduates.count()}')
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph('')  # Пустая строка

        # Статистика
        doc.add_heading('📊 Статистика', level=1)

        employed_count = sum(1 for g in graduates if hasattr(g, 'employment') and g.employment and g.employment.status)
        unemployed_count = graduates.count() - employed_count

        stats_table = doc.add_table(rows=3, cols=2)
        stats_table.style = 'Light List Accent 1'

        stats_data = [
            ('Всего выпускников', str(graduates.count())),
            ('С указанным трудоустройством', str(employed_count)),
            ('Без информации о трудоустройстве', str(unemployed_count)),
        ]

        for i, (label, value) in enumerate(stats_data):
            row = stats_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            row.cells[0].paragraphs[0].runs[0].bold = True

        doc.add_paragraph('')

        # Таблица выпускников
        doc.add_heading('📋 Список выпускников', level=1)

        # Заголовки таблицы
        table = doc.add_table(rows=1, cols=10)
        table.style = 'Light Grid Accent 1'

        headers = ['№', 'ФИО', 'Год', 'Факультет', 'Специальность', 'Email', 'Телефон', 'Статус', 'Работодатель',
                   'Должность']
        hdr_cells = table.rows[0].cells

        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Данные
        for idx, g in enumerate(graduates, 1):
            emp = getattr(g, 'employment', None)

            row_cells = table.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[1].text = g.full_name or '—'
            row_cells[2].text = str(g.graduation_year) if g.graduation_year else '—'
            row_cells[3].text = g.faculty or '—'
            row_cells[4].text = g.specialization or '—'
            row_cells[5].text = g.email or '—'
            row_cells[6].text = g.phone or '—'

            # Статус с цветом
            if emp and emp.status:
                status_para = row_cells[7].paragraphs[0]
                status_run = status_para.add_run(emp.status.name)
                if emp.status.name.lower() == 'трудоустроен':
                    status_run.font.color.rgb = RGBColor(34, 139, 34)
                elif emp.status.name.lower() == 'в поиске':
                    status_run.font.color.rgb = RGBColor(255, 165, 0)
                status_run.bold = True
            else:
                row_cells[7].text = 'Не указан'

            row_cells[8].text = emp.employer.name if emp and emp.employer else '—'
            row_cells[9].text = emp.job_title if emp and emp.job_title else '—'

            # Центрирование номера
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Футер
        doc.add_paragraph('')
        doc.add_paragraph('_' * 100)
        footer = doc.add_paragraph(
            'Документ сформирован автоматически системой учёта трудоустройства выпускников МУ им. С.Ю. Витте'
        )
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.runs[0].font.size = Pt(8)
        footer.runs[0].italic = True

        doc.save(filepath)

    def _export_xlsx(self, graduates, filepath, query):
        """Экспорт в XLSX с форматированием"""
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter

        wb = Workbook()
        ws = wb.active
        ws.title = "Выпускники"

        # Заголовок
        ws.merge_cells('A1:J1')
        title_cell = ws['A1']
        title_cell.value = 'ОТЧЁТ ПО ВЫПУСКНИКАМ'
        title_cell.font = Font(size=16, bold=True, color='FFFFFF')
        title_cell.fill = PatternFill(start_color='940101', end_color='940101', fill_type='solid')
        title_cell.alignment = Alignment(horizontal='center', vertical='center')
        ws.row_dimensions[1].height = 30

        # Метаданные
        ws.merge_cells('A2:J2')
        meta_cell = ws['A2']
        meta_text = f'Дата формирования: {datetime.now().strftime("%d.%m.%Y %H:%M")}'
        if query:
            meta_text += f' | Поисковый запрос: "{query}"'
        meta_text += f' | Найдено: {graduates.count()}'
        meta_cell.value = meta_text
        meta_cell.alignment = Alignment(horizontal='center')
        meta_cell.font = Font(italic=True)

        # Пустая строка
        ws.row_dimensions[3].height = 5

        # Заголовки таблицы
        headers = ['№', 'ФИО', 'Год выпуска', 'Факультет', 'Специальность', 'Email', 'Телефон', 'Статус',
                   'Работодатель', 'Должность']
        header_fill = PatternFill(start_color='D3D3D3', end_color='D3D3D3', fill_type='solid')
        header_font = Font(bold=True, size=11)

        for col_num, header in enumerate(headers, 1):
            cell = ws.cell(row=4, column=col_num, value=header)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            cell.border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )

        ws.row_dimensions[4].height = 30

        # Данные
        for idx, g in enumerate(graduates, 1):
            emp = getattr(g, 'employment', None)
            row_num = idx + 4

            data = [
                idx,
                g.full_name or '—',
                g.graduation_year if g.graduation_year else '—',
                g.faculty or '—',
                g.specialization or '—',
                g.email or '—',
                g.phone or '—',
                emp.status.name if emp and emp.status else 'Не указан',
                emp.employer.name if emp and emp.employer else '—',
                emp.job_title if emp and emp.job_title else '—'
            ]

            for col_num, value in enumerate(data, 1):
                cell = ws.cell(row=row_num, column=col_num, value=value)
                cell.alignment = Alignment(horizontal='left' if col_num > 1 else 'center', vertical='center',
                                           wrap_text=True)
                cell.border = Border(
                    left=Side(style='thin'),
                    right=Side(style='thin'),
                    top=Side(style='thin'),
                    bottom=Side(style='thin')
                )

                # Цвет для статуса
                if col_num == 8 and emp and emp.status:
                    if emp.status.name.lower() == 'трудоустроен':
                        cell.font = Font(color='228B22', bold=True)
                    elif emp.status.name.lower() == 'в поиске':
                        cell.font = Font(color='FFA500', bold=True)
                    else:
                        cell.font = Font(color='DC143C', bold=True)

        # Автоподбор ширины колонок (исправленная версия)
        column_widths = [5, 30, 12, 25, 25, 25, 15, 15, 30, 25]
        for i, width in enumerate(column_widths, 1):
            column_letter = get_column_letter(i)
            ws.column_dimensions[column_letter].width = width

        # Замораживание заголовков
        ws.freeze_panes = 'A5'

        wb.save(filepath)


# ========================
# ОТЧЁТЫ
# ========================

class ReportsView(RoleRequiredMixin, TemplateView):
    """Страница генерации отчётов"""
    allowed_roles = ['manager', 'admin']
    template_name = 'reports.html'


# ========================
# АДМИНКА
# ========================

class AdminUsersView(RoleRequiredMixin, ListView):
    """Список пользователей"""
    allowed_roles = ['admin']
    template_name = 'admin/users.html'
    context_object_name = 'users'
    paginate_by = 25

    def get_queryset(self):
        return User.objects.select_related('role').all()


class AdminCreateUserView(RoleRequiredMixin, View):
    """Создание пользователя админом"""
    allowed_roles = ['admin']
    template_name = 'admin/create_user.html'

    def get(self, request):
        roles = Role.objects.exclude(name='graduate')
        return render(request, self.template_name, {'roles': roles})

    def post(self, request):
        username = request.POST.get('username', '')
        email = request.POST.get('email', '')
        password = request.POST.get('password', '')
        role_id = request.POST.get('role_id', '')

        if User.objects.filter(username=username).exists():
            messages.error(request, "Логин занят")
        elif User.objects.filter(email=email).exists():
            messages.error(request, "Email уже используется")
        else:
            User.objects.create(
                username=username,
                email=email,
                password=make_password(password),
                role_id=int(role_id)
            )
            messages.success(request, "Пользователь создан")
            return redirect('muiv_graduation_system:admin_users')

        roles = Role.objects.exclude(name='graduate')
        return render(request, self.template_name, {'roles': roles})


class PendingRequestsView(RoleRequiredMixin, ListView):
    """Заявки на регистрацию"""
    allowed_roles = ['admin']
    template_name = 'admin/requests.html'
    context_object_name = 'requests'
    paginate_by = 25

    def get_queryset(self):
        return RegistrationRequest.objects.filter(is_approved=False).order_by('-created_at')


class ApproveRequestView(RoleRequiredMixin, View):
    """Одобрение заявки"""
    allowed_roles = ['admin']

    def post(self, request, request_id):
        req = get_object_or_404(RegistrationRequest, id=request_id)

        if req.is_approved:
            messages.warning(request, "Заявка уже одобрена")
            return redirect('muiv_graduation_system:pending_requests')

        role = Role.objects.get(name='manager')
        User.objects.create(
            username=req.username,
            email=req.email,
            password=req.password_hash,
            role=role
        )

        req.is_approved = True
        req.approved_by = request.user
        req.save()

        messages.success(request, f"Менеджер {req.username} успешно добавлен!")
        return redirect('muiv_graduation_system:pending_requests')
